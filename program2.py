import sys
from docx import Document # may need to install this package! 
import xlrd
import xlwt 
import xlsxwriter
#CREATE EXCEL FILE CALLED test.xlsx
workbook = xlsxwriter.Workbook('referenceTable.xlsx') 
worksheet = workbook.add_worksheet()
underline = workbook.add_format({'underline': True})

input_file = sys.argv[1]
freq_file = sys.argv[2]

#-----------------------------------------
# Function(s)

def listToString(li):
	fullString = ""
	for i,unit in enumerate(li):
		if i == 0:
			fullString = unit
			continue
		fullString = fullString + "," + unit
	return fullString

#-----------------------------------------

'''
with open(input_file, 'rb') as f:
	doc = Document(input_file)
	fullText = []
	for para in doc.paragraphs:
		fullText.append(para.text)
	print('\n'.join(fullText))
'''
# Create Excel file to write into

# To read/write into excel files, look at the following snippet:

'''
# Reading an excel file using Python 
import xlrd 
  
# Give the location of the file 
loc = ("path of file") 
  
# To open Workbook 
wb = xlrd.open_workbook(loc) 
sheet = wb.sheet_by_index(0) 
  
# For row 0 and column 0 
sheet.cell_value(0, 0).encode("ascii") 
'''

'''
Underlining parts of cells in Excel

import xlsxwriter
workbook = xlsxwriter.Workbook('test.xlsx')
worksheet = workbook.add_worksheet()
underline = workbook.add_format({'underline': True})
worksheet.write_rich_string(0, 0, 'under', underline, 'line') # 'line' is underlined, 'under' is not.
worksheet.write_rich_string('A2', underline, 'under', 'line') # 'under' is underlined, 'line' is not.
workbook.close()
'''

print("Reading input file...")
# Open the files
unitListFile = open(input_file, 'rb')
ulf = Document(unitListFile)
unitlist = []
for para in ulf.paragraphs:
	unitlist.append(para.text)

print("Reading frequency list into memory...")
# Open and read in the frequency table from the excel file from program 1
freqListFile = xlrd.open_workbook(freq_file)
sheet = freqListFile.sheet_by_index(0)
freqTable = {}
for i in range(0,sheet.nrows):
	freqTable[sheet.cell_value(i,0)] = sheet.cell_value(i,1)

# Create Excel file to write into

# For each line:
i = 0
for units in unitlist:
	# Tokenize the units string
	toks = units.split(",")

	# Search for most frequent unit (down frequency list) -> dubbed B
	topword = None
	topnum = 0
	for token in toks:
		if token in freqTable:
			if freqTable[token] > topnum:
				topnum = freqTable[token]
				topword = token

	# Check if none of the words were found, if so, write the whole line WITHOUT underlines and move on
	if topword == None:
		# Write whole line in alphabetical order
		rest = toks
		rest.sort(key=lambda v: v.upper())
		worksheet.write(i,0,listToString(rest))
		i += 1
		continue

	# Special case: topword is the only one?
	if len(toks) == 1:
		worksheet.write(i,0,topword,underline)
		worksheet.write(i,1,topword)
		i += 1
		continue

	# Remove B from the read line, sort the rest
	rest = toks
	rest.remove(topword)
	rest.sort(key=lambda v: v.upper())
	restString = listToString(rest)

	# Write B,[REST OF LINE] with B underlined into XLSX
	worksheet.write_rich_string(i,0,underline,topword,",",restString)

	# Write B again, normally, into the right column
	worksheet.write(i,1,topword)
	i += 1

# def mainstream(word_list, freq_list):
    



workbook.close()


