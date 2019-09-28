import sys
from docx import Document # may need to install this package! 
import operator
import xlsxwriter

# Select file
input_file = sys.argv[1]

# Output underlined file
output_file = sys.argv[2]

# Create output
document = Document()
document.styles['Normal'].paragraph_format.space_after = 0

# Read in the full file contents into duplist
duplicateFile = open(input_file, 'rb')
ulf = Document(duplicateFile)
duplist = []
for para in ulf.paragraphs:
    duplist.append(para.text)

# Initialize a dictionary for frequency counts
freqCount = {}

print("Counting frequencies and creating underlined redundancy file...")

# Find frequencies and underline redundancies
underline = False
for i,phrase in enumerate(duplist):
    # First element
    if i == 0:
        freqCount[phrase] = 1
        para = document.add_paragraph('')
        run = para.add_run(phrase)
    else:
        if phrase in freqCount:
            freqCount[phrase] += 1
            # Is the line the same as the one above? (redundant)
            if phrase == duplist[i-1]:
                underline = True
        else:
            freqCount[phrase] = 1

        # Write the line in
        para = document.add_paragraph('')
        run = para.add_run(phrase)
        if underline:
            run.underline = True
            underline = False

# Output the underlined file
document.save(output_file)

# Sorting the dictionary by frequency
sortFreq = sorted(freqCount.items(),key = operator.itemgetter(1), reverse=True)

# Excel table output
print("Creating Excel sheet output with frequency counts...")

workbook = xlsxwriter.Workbook('frequencyTable.xlsx') 
worksheet = workbook.add_worksheet()

for (i,(key,val)) in enumerate(sortFreq):
    worksheet.write(i,0,key)
    worksheet.write(i,1,val)

workbook.close()

print("Done!")